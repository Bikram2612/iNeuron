#!/usr/bin/env python
# coding: utf-8

# In[ ]:


get_ipython().set_next_input('1. In what modes should the PdfFileReader() and PdfFileWriter() File objects will be opened');get_ipython().run_line_magic('pinfo', 'opened')


# In[ ]:


read binary (rb)-----> PdfFileREader() and 
write binary (wb)-----> PdfFileWriter()


# In[ ]:


2. From a PdfFileReader object, how do you get a Page object for page 5?

--> by Calling .getNumPages(4)


# In[2]:


get_ipython().system('pip install pyPDF2')


# In[23]:


import PyPDF2 as pdf

pdf_obj_path = "C:/Users/My PC/Desktop/bb.pdf"

file = open(pdf_obj_path, 'rb')
pdf_reader = pdf.PdfFileReader(file)
Total_Page = pdf_reader.getNumPages()
print(" Total No of pages in pdf :", Total_Page, "\n\n")

for i in range(0,2):
    page_obj = pdf_reader.getPage(i)
    print("The text are: \n\n", page_obj.extract_text())


# In[7]:


get_ipython().set_next_input('3. What PdfFileReader variable stores the number of pages in the PDF document');get_ipython().run_line_magic('pinfo', 'document')

---> The total number of pages in the document is stored in the numPages attribute of a PdfFileReader object


# In[ ]:


get_ipython().set_next_input('5. What methods do you use to rotate a page');get_ipython().run_line_magic('pinfo', 'page')

--> rotateClockwise() and rotateCounterClockwise() methods . and passing the no of deg to ratate.


# In[ ]:


get_ipython().set_next_input('6. What is the difference between a Run object and a Paragraph object');get_ipython().run_line_magic('pinfo', 'object')

---> Paragraph object: A document contains multiple paragraphs. A paragraph begins on a new line and 
      contains multiple runs. 
        
---> Run object: Runs are contiguous groups of characters within a paragraph with the same style


# In[ ]:


get_ipython().set_next_input('7. How do you obtain a list of Paragraph objects for a Document object that’s stored in a variable named doc');get_ipython().run_line_magic('pinfo', 'doc')


# In[25]:


get_ipython().system('pip install python-docx')


# In[47]:


import docx
doccument = docx.Document("C:/Users/My PC/Desktop/sample.docx")
for para in doccument.paragraphs:
    fullText = []
    print(para.text)
    fullText.append(para.text)
    '\n'.join(fullText)


# In[ ]:


get_ipython().set_next_input('8. What type of object has bold, underline, italic, strike, and outline variables');get_ipython().run_line_magic('pinfo', 'variables')

--> A Run object has bold, underline,italic,strike and outline variables


# In[ ]:


get_ipython().set_next_input('9. What is the difference between False, True, and None for the bold variable');get_ipython().run_line_magic('pinfo', 'variable')

--> Runs can be further styled using text attributes. Each attribute can be set to one of three values:
True (the attribute is always enabled, no matter what other styles are applied to the run),
False (the attribute is always disabled),
None (defaults to whatever the run’s style is set to)

--> True always makes the Run object bolded and False makes it always not bolded, 
    no matter what the style’s bold setting is. None will make the Run object just use the style’s bold setting


# In[ ]:


get_ipython().set_next_input('10. How do you create a Document object for a new Word document');get_ipython().run_line_magic('pinfo', 'document')

--> by using docx.Document() function.


# In[ ]:


get_ipython().set_next_input(". How do you add a paragraph with the text 'Hello, there!' to a Document object stored in a variable named doc");get_ipython().run_line_magic('pinfo', 'doc')


# In[51]:


import docx
doc = docx.Document()
doc.add_paragraph('Hello there!')
doc.save("C:/Users/My PC/Desktop/doc_new.docx")


# In[ ]:


get_ipython().set_next_input('12. What integers represent the levels of headings available in Word documents');get_ipython().run_line_magic('pinfo', 'documents')

integer from 0 to 4
The arguments to add_heading() are a string of the heading text and an integer from 0 to 4. 
The integer 0 makes the heading the Title style, which is used for the top of the document. 
Integers 1 to 4 are for various heading levels, with 1 being the main heading and 4 the lowest subheading

